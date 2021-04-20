
import boto3
import json
from zipfile import ZipFile
from config import BUCKET
import dynamo as d
import docx
import io
import os
import tempfile

def add_header(doc, key):

    hd = doc.tables[0].cell(0, 0).text
    hd = hd.replace('Pizzaria Abaré Unidade:', 'Pizzaria Abaré Unidade: {}'.format(key['location']))
    hd = hd.replace('Data:', 'Data: {}'.format(key['visit_date']))
    hd = hd.replace('Horário:', 'Horário: {}'.format(key['visit_hour']))

    header = doc.tables[0].cell(0, 0).text = hd
    return doc

def add_bullets(doc, key):

    obs_merge = "\n\n"
    for n, obs in enumerate(key['bullets']):
        obs_merge += "{}. ".format(n)
        obs_merge += obs
        obs_merge += "\n\n"

    p = doc.add_paragraph(text = obs_merge)
    p.paragraph_format.alignment = 0

    return doc

def get_check_session_and_item(row_item):

    session_number = row_item[0]
    item_number = row_item.replace(session_number+'.', '').strip()
    item_number = item_number[:2].strip()
    
    return int(session_number), int(item_number)

def get_items_to_be_checked(doc, key):
    
    unmasked_item = unmask_item_name(key)

    checklist = []
    for table_index, table in enumerate(doc.tables):
        for row in range(50):
            try:
                
                row_item = table.cell(row, 0).text
                row_grade = int(table.cell(row, 3).text)
                session_number, item_number = get_check_session_and_item(row_item)
                
                if row_item in unmasked_item:
                    checklist.append({
                        'name': row_item,
                        'session': session_number,
                        'number': item_number,
                        'grade': row_grade,
                        'table': table_index, 
                        'row':  row
                    })
            except:
                pass
            
    return checklist

def unmask_item_name(key):

    client = boto3.client("s3")
    data = client.get_object(Bucket = BUCKET, Key = "data/item.json")
    items = json.loads(data['Body'].read())
    
    unmask_item = []
    for i in key['irregular']:
        try:
            unmask_item.append(items[i])
        except:
            pass
        
    return unmask_item

def mark_session_total_points(table, row_n, total_points):
    
    t = table.cell(row_n, 3)
    run = t.paragraphs[0].add_run(text = str(total_points))
    run.bold = True
    
    return table

def mark_report_item(table, row_n):

    t = table.cell(row_n, 2)
    run = t.paragraphs[0].add_run(text = "X")
    run.bold = True
    
    return table

def mark_points_per_session(doc, items_to_check):

    # 1. Discount points from sessions marked
    grade_per_session = {1: 93, 2: 43, 3 : 8, 4: 69,5: 16, 6: 59, 7: 43, 8: 25, 9: 30, 10: 20, 11: 18, 12: 23}
    
    for i in items_to_check:
        session = i['session']
        discount_grade = i['grade']
        grade_per_session[session] -= discount_grade

    # 2. Mark total points per session
    for table in doc.tables:
        for row in range(100):
            try:
                text = table.cell(row, 0).text
                if text.find('Total de pontos item') != -1:
                    for s in grade_per_session.keys():
                        if text.find('Total de pontos item {}:'.format(s)) != -1:
                            table = mark_session_total_points(table, row, grade_per_session[s])
            except:
                pass
    return doc



def add_irregular_items(doc, key, T = 447):
    
    # 1. Get items to be checked
    items_to_check = get_items_to_be_checked(doc, key)
    
    # 2. Mark points per session
    doc = mark_points_per_session(doc, items_to_check)
    
    # 3. Mark Items
    
    G = 0
    for item in items_to_check:
        table = mark_report_item(doc.tables[item['table']], item['row'])
        G += int(item['grade'])
        
    # 4. Sum up visit points
        
    P = T - G
    for i in doc.paragraphs:
        if i.text == "Total de pontos atingidos: ":
            p = i.add_run(str(P))
            i.runs[-1].bold = True
    
    # 5. Get classification
    
    rank = P / T
    if rank >= 0.960:
        sts = 'EXCELENTE'
    else: 
        if rank >= 0.90:
            sts = 'BOM'
        else:
            if rank >= 0.51:
                sts = 'REGULAR'
            else:
                sts = 'RUIM'
                
    # 6. Mark visit classification
    
    p = doc.tables[-1].row_cells(0)[0]
    for r in p.paragraphs:
        if r.text.find(sts) != -1:
            r.text = r.text.replace("(  ) {}".format(sts), "(X) {}".format(sts))
            r.runs[-1].bold = True
    
    return doc, {'pontos': P, 'percentual': '{:0.1%}'.format(rank), 'classificacao': sts}

# def add_images(key):

#     client = boto3.client("s3")
#     for img_key in key['png_file']:
        
#         data = client.get_object(Bucket = BUCKET, Key = img_key)
#         img_bytes = io.BytesIO(data['Body'].read())
#         img_bytes.seek(0)

#         p = doc.add_paragraph('')
#         p = p.insert_paragraph_before('')
#         r = p.add_run()
#         r.add_picture(img_bytes, 3000000, 3000000)
#         p.insert_paragraph_before('')
        
#     return doc

def create_report(visit_id):

    key = d.get_report_key(visit_id)
    
    # 1. Temporary dir to save .docx & .zip
    
    temp_dir = tempfile.mkdtemp()
    temp_path = os.path.join(temp_dir)
    
    # 2. Create document file
    
    save_path = "{}/relatorio{}.docx".format(temp_path, key['location'].replace(' ', '').strip())
    
    doc = docx.Document(os.getcwd() + "/doc_template.docx")
    doc = add_header(doc, key)
    doc, result = add_irregular_items(doc, key)
    doc = add_bullets(doc, key)
    doc.save(save_path)
    
    # 3. Create photo zip file
    
    save_zip = "{}/fotos{}.zip".format(temp_path, key['location'].replace(' ', '').strip())
    
    client = boto3.client("s3")
    
    with ZipFile(save_zip, "w") as f:
        for img_key in key['png_file']:
            
            data = client.get_object(Bucket = BUCKET, Key = img_key)
            img_bytes = io.BytesIO(data['Body'].read())
            img_bytes.seek(0)
    
            f.writestr(img_key, img_bytes.getvalue())

    
    # 4. Report metrics
    
    metric = {
        'path': save_path,
        'zip_path': save_zip,
        'classification': result['classificacao'],
        'points': str(result['pontos']),
        'percent_points': result['percentual'],
        'visit_id': visit_id,
        'location': key['location'],
        'item_count': len(key['irregular']),
        'img_count': len(key['png_file']),
        'obs_count': len(key['bullets'])
    }
    
    return metric




